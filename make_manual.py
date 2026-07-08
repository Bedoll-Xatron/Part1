"""MarketFlow 대시보드 활용 가이드 워드 파일 생성"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 페이지 여백 ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── 스타일 헬퍼 ─────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name   = '맑은 고딕'
    run.font.size   = Pt(size)
    run.bold        = bold
    run.italic      = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(31, 78, 121))
    # 하단 테두리
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'■  {text}')
    set_font(run, size=13, bold=True, color=(0, 70, 127))
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'▶  {text}')
    set_font(run, size=11, bold=True, color=(68, 114, 196))
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(level * 0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def tip_box(text):
    """회색 배경 팁 박스 (표 1셀로 구현)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    # 배경색
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EBF3FB')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run('💡 활용 팁  ')
    set_font(run, size=10, bold=True, color=(31, 78, 121))
    run2 = p.add_run(text)
    set_font(run2, size=10, color=(31, 78, 121), italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def workflow_box(steps):
    """워크플로우 화살표 박스"""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F7FB')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run('  →  '.join(steps))
    set_font(run, size=10, bold=True, color=(68, 114, 196))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_kv_table(rows_data):
    """키-값 2열 표"""
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(rows_data):
        # 키 셀
        kc = tbl.cell(i, 0)
        kc.width = Cm(4)
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        set_font(kr, size=10, bold=True, color=(31, 78, 121))
        tcPr = kc._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'D6E4F0')
        tcPr.append(shd)
        # 값 셀
        vc = tbl.cell(i, 1)
        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        set_font(vr, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════
#  표지
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('MarketFlow 대시보드')
set_font(run, size=28, bold=True, color=(31, 78, 121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('메뉴별 상세 활용 가이드')
set_font(run, size=18, color=(68, 114, 196))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run(f'작성일: {datetime.date.today().strftime("%Y년 %m월 %d일")}')
set_font(run, size=11, color=(128, 128, 128))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  목차
# ════════════════════════════════════════════════════════════════
heading1('목  차')
toc = [
    ('1단계. 시장 환경 판단',   ['1. Overview', '2. 시장 고점 탐지', '3. FTD 탐지', '4. 시장 폭 분석']),
    ('2단계. 종목 발굴',        ['5. Best of Best', '6. VCP Signals', '7. VCP 누적 성과',
                                 '8. 종가베팅', '9. 종가베팅 누적', '10. 수급 모멘텀',
                                 '11. 테마 모멘텀', '12. 섹터 로테이션', '13. 역발상', '14. 갭 드리프트']),
    ('3단계. 성과 추적',        ['15. 시그널 사후분석', '16. Backtest Expert']),
    ('4단계. 매매 실행 · 조회', ['17. 포지션 사이징', '18. 종목 검색']),
]
for section_title, items in toc:
    p = doc.add_paragraph()
    run = p.add_run(section_title)
    set_font(run, size=11, bold=True, color=(31, 78, 121))
    for item in items:
        bullet(item, level=1)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  1단계 — 시장 환경 판단
# ════════════════════════════════════════════════════════════════
heading1('1단계. 시장 환경 판단')
body('모든 매매의 전제는 시장 환경 파악입니다. 아래 4개 메뉴를 순서대로 확인한 후 종목 발굴 단계로 넘어가세요.')

# 1. Overview
heading2('1. Overview')
body('KR Market 전체 시장 현황을 한눈에 파악하는 홈 화면입니다.')
heading3('주요 화면 구성')
bullet('KODEX200 차트 — 최근 220일 종가 + MA20 · MA50 · MA200 표시')
bullet('시장 국면 배지 — 지금이 살 때인지 쉬어야 할 때인지를 "매수 우호(RISK_ON) / 중립(NEUTRAL) / 관망(RISK_OFF)" 3단계로 표시합니다. 가장 먼저 확인해야 할 신호등입니다.')
bullet('KOSPI · KOSDAQ 지수 현재가 및 등락률')
bullet('업종별 등락률 — 상위/하위 섹터 색상 구분')
heading3('시장 국면 해석 — 지금 어떤 환경인가?')
body('이동평균선(MA)이란 일정 기간 주가의 평균값을 이은 선입니다. MA200은 최근 200일(약 1년) 평균이고, 지수가 이 선 위에 있으면 "긴 상승 흐름 안에 있다"는 뜻입니다.')
add_kv_table([
    ('매수 우호 국면 (RISK_ON)',
     '지수(KODEX200)가 1년 평균선(MA200) 위에 있고, 단기선(MA20)이 중기선(MA50)보다 높은 상태. '
     '"시장이 우상향 중"이라는 의미 → 모든 전략 적극 활용'),
    ('중립 국면 (NEUTRAL)',
     '상승·하락 조건이 섞여 있는 상태. "확신이 서지 않는 횡보 구간" → 투자금을 줄이고 보수적으로 접근하세요.'),
    ('관망 국면 (RISK_OFF)',
     '지수가 1년 평균선 아래로 내려앉고 단기선도 중기선 아래. "하락세가 뚜렷하다"는 의미 → 신규 매수 자제, 현금 보유 권장'),
])
tip_box('시장 국면이 관망(RISK_OFF)으로 표시되면 아래 메뉴에서 매수 신호가 나와도 투자를 보류하거나 투자금을 절반으로 줄이세요. 좋은 종목도 하락장에서는 함께 떨어집니다.')

# 2. 시장 고점 탐지
heading2('2. 시장 고점 탐지')
body('6개 지표를 복합 점수(0~100)로 산출해 현재 시장이 고점에 얼마나 가까운지 측정합니다.')
heading3('6개 구성 지표')
add_kv_table([
    ('배분일 (25점)',        'KOSPI 거래량 증가 + 0.2% 이상 하락일 수 — 기관 매도 신호'),
    ('선도주 약세 (20점)',   'VCP 신호 종목 중 피벗 고점 하회 비율'),
    ('방어섹터 로테이션 (15점)', '의약품·음식료·통신 vs 반도체·전기전자 상대강도'),
    ('시장 폭 MA60 (15점)', '전체 추적 종목 중 MA60 하회 비율'),
    ('지수 기술 조건 (15점)', 'KODEX200이 MA20·MA50·MA200 몇 개 하회 중인지'),
    ('센티먼트 (10점)',      '역발상 시그널 수 — 많을수록 시장 약세'),
])
heading3('위험 구간 해석')
add_kv_table([
    ('0~20 (Green)',    '정상 — 매수 우호 환경'),
    ('21~40 (Yellow)', '조기 경고 — 신규 진입 신중'),
    ('41~60 (Orange)', '위험 상승 — 포지션 축소 시작'),
    ('61~80 (Red)',    '고점 가능성 높음 — 헤지 고려'),
    ('81~100 (Critical)', '고점 형성 중 — 현금 비중 최대화'),
])
tip_box('점수 40점 이하일 때 Best of Best · VCP 진입을 권장합니다. 60점 초과 시 역발상 전략은 주의, 나머지 전략도 소규모만 운용하세요.')

# 3. FTD 탐지
heading2('3. FTD 탐지 (Follow-Through Day)')
body('William O\'Neil의 FTD 방식으로 시장 저점 이후 반등 확인 신호를 탐지합니다. 조정 후 재진입 타이밍을 잡는 데 사용합니다.')
heading3('FTD 작동 원리 — 시장 저점을 어떻게 확인하나?')
body('시장이 하락하다가 반등할 때, 단순히 한두 번 오른다고 바닥이라고 볼 수 없습니다. O\'Neil의 FTD는 "진짜 반등인지 가짜 반등인지"를 확인하는 3단계 절차입니다.')
add_kv_table([
    ('① 반등 시도 (Rally Attempt)',
     '시장이 하락하다가 처음으로 플러스로 마감하는 날 (전일 대비 +0.3% 이상). '
     '이날부터 날짜를 세기 시작합니다. "드디어 반등 시도가 생겼다"는 신호입니다.'),
    ('② FTD 확인 조건',
     '반등 시도 후 4일 이상 지난 뒤, KOSPI가 +1.5% 이상 오르면서 거래량도 전날보다 많아야 합니다. '
     '"큰 거래량을 동반한 강한 상승" = 기관·외국인이 본격적으로 사들이고 있다는 증거입니다.'),
    ('③ FTD 무효',
     '반등 시도 이후 기존 최저점 아래로 다시 떨어지면 무효. '
     '처음부터 다시 새 저점을 기다려야 합니다.'),
])
heading3('현재 상태별 대응 방법')
add_kv_table([
    ('저점 확인 (CONFIRMED)',
     'FTD 조건 모두 충족 → 시장 저점이 확인됐습니다. 역발상·수급모멘텀 전략을 적극 활용하고 매수 비중을 높이세요.'),
    ('관찰 중 (WATCHING)',
     '반등 시도는 나왔지만 FTD 조건 아직 미충족 → 소규모 테스트 매수만 허용. 본격 진입은 CONFIRMED 이후에'),
    ('무효 (FAILED)',
     '한번 FTD가 됐다가 다시 저점 하향 → 시장이 다시 내려갔습니다. 신규 매수 중단, 보유 포지션 손절 기준 강화'),
    ('신호 없음 (NO_SIGNAL)',
     '현재 상승 추세 중이거나 아직 조정 초기 단계 → 시장 국면(Overview) 기준으로 판단'),
])
tip_box('저점 확인(CONFIRMED) + 시장 고점 탐지 30점 이하 조합이 가장 강력한 매수 신호입니다. 두 조건이 동시에 충족되는 날을 노리세요.')

# 4. 시장 폭 분석
heading2('4. 시장 폭 분석 (Market Breadth)')
body('개별 종목들이 얼마나 골고루 상승하는지를 측정합니다. 지수는 오르지만 소수 대형주만 상승하는 경우를 걸러냅니다.')
heading3('주요 지표')
add_kv_table([
    ('MA20/60/120/200 상위 %', '각 이동평균 위에 있는 종목 비율 — 숫자가 높을수록 건강한 상승'),
    ('등락 종목 수 (A/D)',     '상승 종목 수 ÷ 하락 종목 수 — 1.0 초과면 상승 우세'),
    ('Naver 실시간 A/D',       'KOSPI·KOSDAQ 전체 상장 종목 기준 등락 수'),
    ('52주 신고/신저 비율',    '신고가 종목 ÷ 신저가 종목 — 2.0 이상이면 확장 국면'),
])
heading3('해석 기준')
add_kv_table([
    ('MA20 > 70%',  '강한 단기 상승장 → 모멘텀 전략 전력 투구'),
    ('MA60 < 40%',  '중기 추세 약화 → 포지션 규모 축소'),
    ('A/D < 0.8',   '하락 종목 우세 → 분산 매도, 신규 진입 자제'),
    ('신고/신저 > 2.0', '확장 국면 → 적극 매수 가능'),
    ('신고/신저 < 0.5', '위축 국면 → 현금 비중 확대'),
])
tip_box('시장 고점 탐지 점수와 시장 폭 분석을 함께 보세요. 고점 탐지 낮음 + MA60 상위 60% 이상 = 최적 진입 환경입니다.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  2단계 — 종목 발굴
# ════════════════════════════════════════════════════════════════
heading1('2단계. 종목 발굴')
body('시장 국면이 매수 우호(RISK_ON)일 때 아래 메뉴에서 종목을 선별합니다. 시장 국면이 관망(RISK_OFF)이거나 고점 탐지 점수가 60점 이상이면 좋아 보이는 종목이 있더라도 매수를 보류하세요. 좋은 종목을 찾는 것보다 시장 환경 판단이 먼저입니다.')

# 5. Best of Best
heading2('5. Best of Best')
body('6개 전략(수급·테마·섹터·역발상) 시그널을 통합해 종합 점수 기준 상위 30종목을 보여줍니다.')
heading3('활용 방법')
bullet('전략을 고르기 어려울 때 첫 번째로 확인하는 통합 뷰')
bullet('종합 점수(0~10) 높은 순 정렬 — 7점 이상은 적극 고려')
bullet('출처 전략 배지로 어떤 로직에서 발굴됐는지 확인')
bullet('순위 1~3위(금·은·동 메달 표시)는 당일 최우선 관심 종목')
heading3('주요 컬럼')
add_kv_table([
    ('마켓 배지',    'KP(KOSPI) / KQ(KOSDAQ) 원형 배지'),
    ('종합 Score',  '0~10점 — 모든 전략 지표를 정규화한 통합 점수'),
    ('출처 전략',   '수급·테마·섹터·역발상 중 어느 전략에서 발굴됐는지'),
    ('주요 지표',   '출처 전략별 핵심 수치 (수급 억원 / 테마명 / RS 점수 / 반전 확률)'),
])
tip_box('종목 클릭 시 차트 모달이 열립니다. 캔들 + 거래량 차트로 패턴을 직접 확인하세요.')

# 6. VCP Signals
heading2('6. VCP Signals')
body('Mark Minervini의 Volatility Contraction Pattern(변동성 수축 패턴)을 탐지합니다. 상승 추세에서 변동성이 점진적으로 수축된 후 돌파하는 종목을 포착합니다.')
heading3('VCP 패턴 이해 — 왜 변동성이 줄어드는 게 좋은 신호인가?')
body('주가가 상승 후 조정받을 때, 처음엔 크게 흔들리다가 점점 조용해지는 구간이 생깁니다. 이것이 VCP입니다. 마치 용수철을 눌렀다 놓기 직전처럼 에너지가 응축된 상태입니다.')
bullet('C1 → C2 → C3: 주가가 조정받는 폭이 점점 줄어들어야 합니다. 예) 처음엔 -10%, 다음엔 -6%, 마지막엔 -3% 식으로 수축')
bullet('수축 비율: 각 구간의 조정폭을 비교한 값. 수축이 잘 될수록 고품질 패턴으로 점수가 높아집니다.')
bullet('돌파 기준가(Pivot High): VCP 패턴에서 가장 최근 고점. 이 가격을 거래량 급증과 함께 돌파하면 바로 매수 신호입니다.')
heading3('등급별 의미')
add_kv_table([
    ('A등급', '수축 비율 최대, 고점 하강 + 저점 상승 모두 충족, 장기 MA 위에 위치 → 적극 매수'),
    ('B등급', '수축 비율 양호, MA20 위에 위치 → 관심 유지, 돌파 확인 후 진입'),
    ('C등급', 'MA60 위에 위치하나 패턴 완성도 낮음 → 소규모 관찰'),
    ('D등급', '최소 조건 충족 → 패스 권장'),
])
heading3('언제 사는가? — 진입 조건')
bullet('돌파 기준가(Pivot High) 돌파 + 당일 거래량이 20일 평균의 1.5배 이상 → 이 두 조건이 동시에 충족되는 날에 매수')
bullet('돌파 당일 양봉이 크고 길수록(장대 양봉) 더욱 강한 신호입니다.')
bullet('손절 기준: 돌파 기준가보다 7~8% 아래로 내려가면 손절. 또는 패턴의 최근 저점이 무너지는 경우')
tip_box('VCP 등급 A + 시장 국면 매수 우호(RISK_ON) + FTD 저점 확인(CONFIRMED) — 이 세 가지가 동시에 맞으면 가장 좋은 매수 환경입니다.')

# 7. VCP 누적 성과
heading2('7. VCP 누적 성과')
body('과거 VCP 시그널 전체 이력과 수익률 결과를 추적합니다.')
heading3('주요 통계')
add_kv_table([
    ('승률',         '종가 기준 목표가 도달 비율'),
    ('평균 수익률',   '청산된 포지션 평균 ROI %'),
    ('Profit Factor', '총 수익 ÷ 총 손실 — 1.5 이상이면 양호'),
    ('등급별 통계',   'A/B/C 등급별 승률 및 평균 수익 분리 확인'),
])
heading3('활용 방법')
bullet('등급별 성과 차이 확인 → 하위 등급 필터링 기준 설정')
bullet('기간별 승률 변화로 이 전략이 다양한 시장 환경에서도 꾸준히 통하는지 판단')
bullet('결과가 좋지 않은 기간 → 시장 환경과 연관성 분석')

# 8. 종가베팅
heading2('8. 종가베팅')
body('장 마감 직전(14:50~15:20) 특정 조건을 충족한 종목을 매수해 다음 날 갭 상승 또는 단기 모멘텀을 노리는 전략입니다.')
heading3('시그널 조건')
bullet('당일 급등 + 거래량 급증 조건 충족')
bullet('스코어링 시스템(15점 만점)으로 점수화 — 9점 이상 A등급')
bullet('외국인·기관 수급 점수 포함')
heading3('매매 방법')
add_kv_table([
    ('진입 시점', '당일 장 마감 10~30분 전 시그널 확인 후 종가 근처 매수'),
    ('목표가',   '진입가 +9~15% (등급에 따라 다름)'),
    ('손절가',   '진입가 -5%'),
    ('보유 기간', '1~5일 단기'),
])
tip_box('종가베팅은 당일 시장 분위기가 강할 때 효과가 큽니다. Overview에서 KOSPI 당일 등락이 +0.5% 이상이면 신뢰도 상승합니다.')

# 9. 종가베팅 누적
heading2('9. 종가베팅 누적')
body('종가베팅 과거 시그널 전체 이력을 목표가/손절가 도달 기준으로 추적합니다.')
heading3('결과 분류')
add_kv_table([
    ('TARGET_HIT', '목표가 도달 — 수익 실현'),
    ('STOP_HIT',   '손절가 이탈 — 손실 확정'),
    ('OPEN',       '아직 미결 — 현재 보유 중'),
])
heading3('활용 방법')
bullet('등급별 TARGET_HIT 비율 비교 → A등급 집중 여부 판단')
bullet('특정 기간 STOP_HIT 급증 시 해당 시장 환경 재분석')
bullet('Backtest Expert와 함께 보면 월별 성과 확인 가능')

# 10. 수급 모멘텀
heading2('10. 수급 모멘텀')
body('외국인·기관의 순매수 금액이 집중되는 종목을 스마트머니 흐름으로 포착합니다.')
heading3('주요 지표')
add_kv_table([
    ('Foreign Flow',      '5일 외국인 순매수 금액 (억원) — 양수면 매수 우세'),
    ('Institution Flow',  '5일 기관 순매수 금액 (억원)'),
    ('Volume Ratio',      '거래대금 대비 수급 강도'),
    ('Signal Strength',   'Strong / Moderate / Weak — 외국인+기관 동시 매수 여부'),
])
heading3('활용 방법')
bullet('Strong 등급 종목 우선 — 외국인+기관 동시 순매수는 가장 강한 신호')
bullet('Foreign Flow + Institution Flow 모두 양수인 종목만 선별')
bullet('VCP 패턴과 겹치는 종목이 있으면 최우선 후보')
tip_box('수급 모멘텀 종목이 Best of Best 상위권에도 있다면 복수 전략이 동시에 선택한 종목입니다. 가장 강한 신호로 간주하세요.')

# 11. 테마 모멘텀
heading2('11. 테마 모멘텀')
body('뉴스·SNS에서 급부상하는 테마와 연관 종목을 점수화합니다. 테마 사이클 초기에 진입하는 전략입니다.')
heading3('주요 지표')
add_kv_table([
    ('테마명',           '종목이 속한 주요 테마 (예: AI·2차전지·바이오)'),
    ('뉴스 센티먼트',   '-1~+1 점수 — 양수일수록 긍정 뉴스 비중 높음'),
    ('SNS 모멘텀',      '0~100 점수 — 소셜미디어 언급량 급증 강도'),
    ('내러티브 점수',   '테마의 지속 가능성 종합 점수'),
])
heading3('활용 방법')
bullet('뉴스 센티먼트 +0.5 이상 + SNS 모멘텀 70 이상 = 강한 테마 부각 신호')
bullet('같은 테마 내 여러 종목이 동시에 등장하면 테마 주도주 파악')
bullet('테마 수명은 짧으므로 진입 후 빠른 익절 전략 권장')
tip_box('테마 모멘텀은 변동성이 크므로 포지션 사이징에서 리스크 비율을 평소보다 낮게(0.5%) 설정하세요.')

# 12. 섹터 로테이션
heading2('12. 섹터 로테이션')
body('업종 순환 사이클에서 현재 주도 업종을 파악하고 해당 업종 내 선도 종목을 발굴합니다.')
heading3('주요 지표')
add_kv_table([
    ('섹터명',         '종목이 속한 업종'),
    ('상대강도 (RS)',  '전체 시장 대비 해당 업종의 상대 성과 점수'),
    ('로테이션 국면', 'Leading / Weakening / Lagging / Improving 4분면'),
])
heading3('로테이션 국면 해석 — 어떤 업종에 올라타야 하나?')
body('상대강도(RS)란 전체 시장 대비 "이 업종이 얼마나 잘 오르고 있는지"를 나타내는 점수입니다. RS가 높다 = 남들보다 강하게 상승 중.')
add_kv_table([
    ('선도 업종 (Leading)',
     '상대강도 높음 + 계속 상승 중 → 지금 시장을 이끄는 주도 업종. 매수에 집중하세요. 가장 먼저 선택해야 할 업종입니다.'),
    ('개선 업종 (Improving)',
     '아직 상대강도는 낮지만 점점 좋아지는 추세 → "곧 선도 업종으로 올라설 가능성"이 있는 업종. 조금 일찍 선점하는 전략입니다.'),
    ('약화 업종 (Weakening)',
     '상대강도는 높지만 모멘텀이 꺾이는 중 → 한때 주도 업종이었으나 이제 빠져나올 시점. 보유 중이라면 익절 준비하세요.'),
    ('뒤처진 업종 (Lagging)',
     '상대강도 낮음 + 계속 하락 → 시장에서 소외된 업종. 이유 없이 진입하지 마세요.'),
])
tip_box('Leading 섹터 종목 + VCP A등급 조합이 섹터 로테이션 전략의 최적 조합입니다.')

# 13. 역발상
heading2('13. 역발상 (Contrarian Reversal)')
body('과매도 구간에 진입한 종목 중 반전 가능성이 높은 종목을 스크리닝합니다. FTD 탐지와 연동하여 사용하면 효과적입니다.')
heading3('주요 지표')
add_kv_table([
    ('과매도 점수',    '0~10 — RSI·볼린저 밴드·이격도 기반 과매도 정도'),
    ('반전 확률',      '0~100% — 유사 패턴의 과거 반전 성공률'),
    ('지지 레벨',      '핵심 지지가격 — 이 가격 하향 이탈 시 손절'),
])
heading3('활용 방법')
bullet('FTD CONFIRMED 상태에서 역발상 시그널 활용도가 가장 높음')
bullet('반전 확률 70% 이상 + 지지 레벨 근처에서 진입')
bullet('손절: 지지 레벨 -3% 이탈 시 즉시 청산')
bullet('목표: 직전 고점 또는 MA60까지')
tip_box('역발상 전략은 시장 고점 탐지 점수가 낮을 때(30점 이하) 효과적입니다. 고점 근방에서는 과매도 반등이 약합니다.')

# 14. 갭 드리프트
heading2('14. 갭 드리프트 (PEAD Proxy)')
body('대형 갭 상승 이벤트 이후 추세가 지속되는 종목을 5요소 점수로 스크리닝합니다. 실적 발표 등 대형 이벤트 후 드리프트를 포착합니다.')
heading3('5요소 점수 (각 0~2점, 합계 0~10점)')
add_kv_table([
    ('갭 크기',    '갭 상승률 — 10% 이상 2점, 6% 이상 1.5점'),
    ('거래량 배수', '갭 당일 거래량 ÷ 20일 평균 — 4배 이상 2점'),
    ('MA 위치',    '현재가가 MA20·MA60 위에 있는지'),
    ('갭 이전 추세', '갭 발생 전 20일간 상승 추세 여부'),
    ('드리프트',   '갭 이후 가격 방향 — 상승 지속 시 2점'),
])
heading3('필터 조정 방법')
bullet('최소 갭 %: 기본 3% — 낮추면 더 많은 종목, 높이면 강한 이벤트만')
bullet('최소 거래량 배수: 기본 1.5 — 높이면 거래량 폭발한 종목만')
bullet('탐색 기간: 기본 60일 — 최근 이벤트만 보려면 20~30일로 축소')
tip_box('점수 7점 이상 + 드리프트 양수(갭 이후 상승 지속)인 종목이 최우선 후보입니다.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  3단계 — 성과 추적
# ════════════════════════════════════════════════════════════════
heading1('3단계. 성과 추적')
body('매매 결과를 데이터로 검증하고 전략별 강약점을 파악합니다.')

# 15. 시그널 사후분석
heading2('15. 시그널 사후분석 (Signal Postmortem)')
body('6개 전략별 과거 시그널의 5일·20일 실제 수익률을 자동 계산해 전략별 성과를 비교합니다.')
heading3('주요 지표')
add_kv_table([
    ('5일 승률',   '시그널 발생 후 5영업일 뒤 수익률이 양수인 비율'),
    ('20일 승률',  '시그널 발생 후 20영업일 뒤 수익률이 양수인 비율'),
    ('5일 평균',   '5일 후 평균 수익률 — 단기 알파 측정'),
    ('20일 평균',  '20일 후 평균 수익률 — 중기 알파 측정'),
])
heading3('활용 방법')
bullet('전략 카드 클릭 → 최근 5건 상세 결과 확인')
bullet('승률이 낮은 전략은 진입 비중 줄이고 높은 전략에 집중')
bullet('최소 20건 이상 누적 후 통계적 유의성 판단')
tip_box('시그널 사후분석에서 특정 전략의 승률이 지속 하락 중이면 시장 환경 변화 신호일 수 있습니다. Backtest Expert의 롤링 승률과 함께 확인하세요.')

# 16. Backtest Expert
heading2('16. Backtest Expert')
body('전략별 상세 백테스트 결과를 수익 곡선·최대 낙폭·월별 성과 등 다차원으로 분석합니다.')
heading3('주요 분석 항목')
add_kv_table([
    ('5일/20일 승률',    '기간별 방향 예측 정확도'),
    ('평균 수익',        '20일 기준 per trade 평균 수익률'),
    ('Profit Factor',    '총 수익 ÷ 총 손실 — 1.5 이상 양호, 2.0 이상 우수'),
    ('최대 낙폭 (MDD)',  '수익 곡선 최고점 → 최저점 하락 폭 % (1% 리스크 복리 기준)'),
    ('최대 연속 손실',   '연속 손실 최대 횟수 — 심리적 버팀목 계획에 활용'),
    ('최근 30건 승률',   '누적 전체 승률 대비 최근 추세 비교'),
    ('수익 곡선',        '최근 60거래 복리 누적 수익 스파크라인'),
    ('월별 히트맵',      '월별 평균 수익률 색상 — 녹색(+), 적색(-) 직관적 파악'),
])
heading3('활용 방법')
bullet('전략 탭을 선택해 각 전략 상세 확인')
bullet('PF ≥ 1.5 + MDD ≤ 15% + 롤링 승률 ≥ 50% = 현재 유효한 전략')
bullet('월별 히트맵에서 계절성(특정 월 부진) 파악 → 해당 월 진입 축소')
bullet('수익 곡선이 우하향 중이면 전략 재검토 신호')
tip_box('모든 전략의 PF와 MDD를 비교 테이블에서 한눈에 확인하고, 현재 시장에서 가장 강한 전략에 비중을 집중하세요.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  4단계 — 매매 실행 · 조회
# ════════════════════════════════════════════════════════════════
heading1('4단계. 매매 실행 · 조회')

# 17. 포지션 사이징
heading2('17. 포지션 사이징 (Position Sizer)')
body('종목 선정 후 매수 수량과 리스크 금액을 계산합니다. 감으로 수량을 정하지 말고 반드시 이 계산기를 사용하세요.')
heading3('3가지 계산 방식')
add_kv_table([
    ('Fixed Fractional', '자본의 고정 % (기본 1%)를 리스크로 사용 — 가장 단순하고 범용적'),
    ('ATR 기반',          'ATR(평균 진폭) × 배수로 손절 거리 결정 — 변동성 자동 반영'),
    ('Kelly Criterion',  '과거 승률·손익비로 수학적 최적 비중 계산 — Half Kelly 사용 권장'),
])
heading3('입력값 및 결과')
add_kv_table([
    ('총 자본금',         '전체 투자 가능 금액'),
    ('진입가',            '매수 예정 가격'),
    ('손절가 / ATR',      '손절 기준가 또는 ATR 값'),
    ('리스크 비율 (%)',   '1거래 허용 손실 비율 — 기본 1%, 최대 2% 권장'),
    ('매수 수량',         '계산된 최적 주수'),
    ('포트폴리오 리스크', '전체 자본 대비 실제 손실 가능 금액 비율 — 2% 초과 시 경고'),
])
heading3('리스크 관리 원칙')
bullet('거래당 리스크 1% 원칙 — 연속 100번 손절해도 자본의 63%는 보존')
bullet('동시 오픈 포지션 총 리스크 6~8% 이하 유지')
bullet('포지션 사이징 결과에서 리스크 경고(빨간색) 표시 시 수량 줄이기')
tip_box('Kelly Criterion은 과거 승률이 55% 이상이고 손익비 1.5 이상일 때 효과적입니다. 초보자는 Fixed Fractional 1%를 권장합니다.')

# 18. 종목 검색
heading2('18. 종목 검색')
body('종목 코드 또는 이름으로 개별 종목을 검색하고 차트, 지표, AI 요약을 확인합니다.')
heading3('검색 결과 정보')
add_kv_table([
    ('현재가·등락률',     '실시간 현재가 및 전일 대비 등락률'),
    ('시가총액·PER·PBR', '기본 밸류에이션 지표'),
    ('VCP·종가베팅 시그널', '현재 해당 종목이 각 전략 시그널에 포함되어 있는지 여부'),
    ('차트 히스토리',     '일봉 캔들 + 거래량 차트 (최근 252일)'),
    ('AI 요약',           '뉴스·공시 기반 종목 현황·전망·리스크 요인 AI 분석'),
])
heading3('활용 방법')
bullet('매수 전 최종 확인: VCP/종가베팅 시그널 여부 + 차트 패턴 육안 확인')
bullet('AI 요약에서 리스크 요인 확인 후 특이사항 없으면 진입')
bullet('다른 메뉴에서 종목명 클릭 시 자동으로 이 화면으로 연결')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  통합 워크플로우
# ════════════════════════════════════════════════════════════════
heading1('통합 워크플로우 — 하루 루틴')

heading2('장 시작 전 (오전 8:30~9:00)')
workflow_box(['Overview\n시장 국면 확인', '시장 고점 탐지\n점수 확인', 'FTD 탐지\n상태 확인', '시장 폭\nA/D 확인'])
bullet('관망 국면(RISK_OFF) + 고점 탐지 60점 이상 → 오늘은 관망. 어떤 좋은 신호가 나와도 매수 금지')
bullet('매수 우호 국면(RISK_ON) + 고점 탐지 30점 이하 + FTD 저점 확인(CONFIRMED) → 적극 매수 환경. 종목 선정에 집중')

heading2('종목 선정 (오전 9:00~9:30)')
workflow_box(['Best of Best\n상위 확인', 'VCP 7점↑\n필터', '수급 모멘텀\nStrong 필터', '포지션 사이징\n수량 계산'])
bullet('Best of Best 상위 10종목 중 VCP 또는 수급 모멘텀과 겹치는 종목 최우선')
bullet('종목 검색에서 차트·AI 요약으로 최종 확인')

heading2('장중 모니터링')
workflow_box(['종가베팅\n조건 확인', '테마 모멘텀\n급등 테마', '갭 드리프트\n이벤트 종목'])
bullet('14:50~15:10 종가베팅 시그널 확인 → 당일 종가 매수 결정')

heading2('주간 성과 검토 (주말)')
workflow_box(['시그널 사후분석\n승률 점검', 'Backtest Expert\nPF·MDD 확인', '전략 비중\n재조정'])
bullet('승률 하락 전략 비중 축소, 상승 전략 비중 확대')
bullet('월별 히트맵에서 부진 기간 패턴 파악')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  주요 용어 정리
# ════════════════════════════════════════════════════════════════
heading1('주요 용어 정리')
add_kv_table([
    ('VCP (변동성 수축 패턴)',
     '주가 조정 폭이 점점 줄어드는 패턴. 마크 미네르비니가 정립한 모멘텀 매매 기법. '
     '변동성이 수축될수록 돌파 이후 큰 상승이 나올 가능성이 높습니다.'),
    ('FTD (시장 저점 확인)',
     'Follow-Through Day — 시장이 바닥을 찍고 진짜 반등을 시작했음을 확인하는 신호. '
     '반등 시도 후 4일 이상 지나 KOSPI가 거래량 급증을 동반한 +1.5% 이상 상승일 때 확인됩니다.'),
    ('배분일 (Distribution Day)',
     '지수는 하락했는데 거래량은 늘어난 날. "기관·외국인이 주식을 팔고 있다"는 의심이 드는 날. '
     '짧은 기간에 이런 날이 많이 쌓이면 시장 고점 신호입니다.'),
    ('반등 시도 (Rally Attempt)',
     '시장이 하락하다가 처음으로 플러스 마감하는 날. FTD를 기다리는 출발점. '
     '이날부터 날짜를 세어 4일 이후 FTD 조건을 확인합니다.'),
    ('돌파 기준가 (Pivot High)',
     'VCP 패턴에서 가장 최근 고점. 이 가격을 거래량 급증과 함께 돌파하는 날이 매수 진입 시점입니다.'),
    ('이동평균선 (MA)',
     '일정 기간 주가의 평균을 이은 선. MA20 = 최근 20일(약 1개월), MA50 = 약 2.5개월, '
     'MA200 = 약 1년 평균. 주가가 MA200 위에 있으면 장기 상승 흐름 안에 있다는 뜻입니다.'),
    ('매수 우호 국면 (RISK_ON)',
     '주식을 사도 좋은 환경. 지수가 주요 이동평균선 위에 있고 상승 모멘텀이 살아있는 상태. '
     '"위험자산(주식)을 선호하는 시장 분위기"를 뜻합니다.'),
    ('관망 국면 (RISK_OFF)',
     '주식 매수를 자제해야 할 환경. 지수가 이동평균선 아래로 내려간 상태. '
     '"안전자산(예금·채권)을 선호하는 시장 분위기"를 뜻합니다.'),
    ('등락 비율 (A/D Ratio)',
     'Advance/Decline Ratio — 오늘 오른 종목 수 ÷ 내린 종목 수. '
     '1.0 초과면 상승 종목이 더 많다는 뜻. 지수는 올라도 A/D가 낮으면 소수 대형주만 오른 상황입니다.'),
    ('손익비 (Profit Factor, PF)',
     '총 수익금 ÷ 총 손실금. 예를 들어 PF=2.0이면 1원을 잃을 때 2원을 버는 전략. '
     '1.5 이상이면 양호, 2.0 이상이면 우수한 전략입니다.'),
    ('최대 낙폭 (MDD)',
     'Maximum Drawdown — 수익 곡선이 최고점에서 최저점까지 떨어진 폭(%). '
     '예) MDD -20%면 한 때 계좌가 20% 줄어든 적이 있다는 뜻. 낮을수록 안정적인 전략입니다.'),
    ('하프 켈리 (Half Kelly)',
     '켈리 공식으로 계산된 최적 투자 비율의 절반만 사용하는 방법. '
     '켈리 원값은 너무 공격적이라 실전에서는 절반이 안전합니다.'),
    ('평균 변동폭 (ATR)',
     'Average True Range — 최근 N일 동안 하루 주가가 평균 얼마나 움직였는지. '
     'ATR이 크면 변동성이 큰 종목, 작으면 안정적인 종목입니다. 손절 거리를 정할 때 활용합니다.'),
    ('갭 드리프트 (PEAD)',
     '큰 갭 상승(주로 실적 발표 등 이벤트) 이후 주가가 계속 오르는 현상. '
     '이벤트 발표 후 한 달 이상 강세가 이어지는 종목을 노리는 전략입니다.'),
    ('최근 승률 (Rolling Win Rate)',
     '전체 누적 승률이 아니라 "최근 N건"만 따진 승률. '
     '전략이 지금도 유효한지 확인하는 데 쓰입니다. 누적 승률은 좋아도 최근 승률이 낮으면 전략 재검토 필요'),
])

# ════════════════════════════════════════════════════════════════
#  저장
# ════════════════════════════════════════════════════════════════
out_path = 'D:/INFORUN/HoDoo/Part7/MarketFlow_활용가이드.docx'
doc.save(out_path)
print(f'저장 완료: {out_path}')
